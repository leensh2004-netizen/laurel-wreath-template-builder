from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_paragraph_rtl(paragraph) -> None:
    """
    Makes paragraph RTL for Arabic text.
    """
    p_pr = paragraph._p.get_or_add_pPr()

    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)

    bidi.set(qn("w:val"), "1")


def replace_paragraph_text_keep_style(paragraph, new_text: str) -> None:
    """
    Replaces paragraph text but keeps the first run style.
    """
    new_text = str(new_text or "").strip()

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text

    for run in paragraph.runs[1:]:
        run.text = ""


def clean_text(value) -> str:
    return " ".join(str(value or "").split())


def update_cover_page(doc, form: dict) -> None:
    """
    Updates the first page / cover page only.

    It replaces the first 4 non-empty paragraphs:
    1. Company name
    2. Legal form
    3. Financial statements title
    4. Cover date
    """

    company_name = str(form.get("company_name", "") or "").strip()
    legal_form = str(form.get("legal_form", "") or "").strip()
    statements_title = str(form.get("statements_title", "") or "").strip()
    cover_date = str(form.get("cover_date", "") or "").strip()

    if not legal_form:
        legal_form = "(ذات مسؤولية محدودة)"

    if not statements_title:
        statements_title = "القوائم الماليــة المنفصلة للسنة المالية المنتهية في"

    if not cover_date:
        cover_date = str(form.get("financial_year", "") or "").strip()

    cover_values = [
        company_name,
        legal_form,
        statements_title,
        cover_date,
    ]

    first_non_empty_paragraphs = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            first_non_empty_paragraphs.append(paragraph)

        if len(first_non_empty_paragraphs) == 4:
            break

    for paragraph, value in zip(first_non_empty_paragraphs, cover_values):
        if value:
            replace_paragraph_text_keep_style(paragraph, value)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_rtl(paragraph)

    update_document_headers(doc, form)


def update_document_headers(doc, form: dict) -> None:
    """
    Updates repeated document headers:
    - Company name
    - Legal form
    - Date line in headers
    """

    company_name = str(form.get("company_name", "") or "").strip()
    legal_form = str(form.get("legal_form", "") or "").strip()
    cover_date = str(form.get("cover_date", "") or "").strip()

    if not legal_form:
        legal_form = "(ذات مسؤولية محدودة)"

    if not cover_date:
        cover_date = str(form.get("financial_year", "") or "").strip()

    header_date_line = ""
    if cover_date:
        header_date_line = f"للسنة المالية المنتهية في {cover_date}"

    for section in doc.sections:
        header_containers = [
            section.header,
            section.first_page_header,
            section.even_page_header,
        ]

        for header in header_containers:
            update_header_container(
                header,
                company_name=company_name,
                legal_form=legal_form,
                header_date_line=header_date_line,
            )


def update_header_container(header, company_name: str, legal_form: str, header_date_line: str) -> None:
    """
    Updates paragraphs inside one Word header.
    Also checks tables inside headers.
    """

    for paragraph in header.paragraphs:
        update_header_paragraph(
            paragraph,
            company_name=company_name,
            legal_form=legal_form,
            header_date_line=header_date_line,
        )

    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    update_header_paragraph(
                        paragraph,
                        company_name=company_name,
                        legal_form=legal_form,
                        header_date_line=header_date_line,
                    )


def update_header_paragraph(paragraph, company_name: str, legal_form: str, header_date_line: str) -> None:
    text = clean_text(paragraph.text)

    if not text:
        return

    # Company name line
    if company_name and text.startswith("شركة"):
        replace_paragraph_text_keep_style(paragraph, company_name)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_rtl(paragraph)
        return

    # Legal form line
    if legal_form and "ذات مسؤولية محدودة" in text:
        replace_paragraph_text_keep_style(paragraph, legal_form)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_rtl(paragraph)
        return

    # Header date line
    if header_date_line and "للسنة المالية المنتهية في" in text:
        replace_paragraph_text_keep_style(paragraph, header_date_line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_rtl(paragraph)
        return
