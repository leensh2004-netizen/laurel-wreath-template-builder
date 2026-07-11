from __future__ import annotations

import copy
import io
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from docx.shared import Inches
from cover_page_editor import update_cover_page
from index_editor import update_page_two_and_index

TEMPLATE_PATH = Path(__file__).parent / "templates" / "نموذج الايضاحات.docx"


@dataclass
class SectionDef:
    key: str
    title: str
    pattern: str
    group: str
    default: bool = True


POLICY_SECTIONS: List[SectionDef] = [
    SectionDef("policy_basis", "2-1 أسس الإعداد", r"2\s*-\s*1\s*.*[اأإ]سس|2\s*-\s*1\s*.*ال[اإ]عداد", "السياسات المحاسبية"),
    SectionDef("policy_currency", "2-3 العملة التشغيلية وعملة عرض القوائم المالية", r"2\s*-\s*3\s+العملة", "السياسات المحاسبية"),
    SectionDef("policy_foreign_currency", "2-4 العمليات والأرصدة", r"2\s*-\s*4\s+العمليات", "السياسات المحاسبية"),
    SectionDef("policy_capital", "2-5 رأس المال", r"2\s*-\s*5\s+ر[أا]س", "السياسات المحاسبية"),
    SectionDef("policy_cash", "2-6 النقد وما في حكمه", r"2\s*-\s*6\s+النقد", "السياسات المحاسبية"),
    SectionDef("policy_receivables", "2-7 الذمم المدينة التجارية", r"2\s*-\s*7\s+الذمم", "السياسات المحاسبية"),
    SectionDef("policy_inventory", "2-8 المخزون", r"2\s*-\s*8\s+المخزون", "السياسات المحاسبية"),
    SectionDef("policy_related_parties", "2-9 الأطراف ذات علاقة", r"2\s*-\s*9\s+ال[اأ]طراف", "السياسات المحاسبية"),
    SectionDef("policy_ppe", "2-10 الممتلكات والمعدات", r"2\s*-\s*10\s+الممتلكات", "السياسات المحاسبية"),
    SectionDef("policy_intangibles", "2-11 الموجودات غير الملموسة", r"2\s*-\s*11\s+الموجودات غير الملموسة", "السياسات المحاسبية"),
    SectionDef("policy_impairment_nonfinancial", "2-12 التدني في قيمة الموجودات غير المالية", r"2\s*-\s*12\s+.*التدني", "السياسات المحاسبية"),
    SectionDef("policy_offset", "2-13 التقاص", r"2\s*-\s*13\s+.*التق", "السياسات المحاسبية"),
    SectionDef("policy_payables", "2-14 الذمم الدائنة والأرصدة الدائنة الأخرى", r"2\s*-\s*14\s+.*الذمم الدائنة", "السياسات المحاسبية"),
    SectionDef("policy_tax", "2-15 ضريبة الدخل", r"2\s*-\s*15\s+ضريبة", "السياسات المحاسبية"),
    SectionDef("policy_loans", "2-16 القروض", r"2\s*-\s*16\s+القروض", "السياسات المحاسبية"),
    SectionDef("policy_borrowing", "2-17 تكاليف الاقتراض", r"2\s*-\s*17\s+تكاليف", "السياسات المحاسبية"),
    SectionDef("policy_revenue", "2-18 تحقق الإيرادات", r"2\s*-\s*18\s+تحقق", "السياسات المحاسبية"),
    SectionDef("policy_deferred_revenue", "2-19 الإيرادات المؤجلة", r"2\s*-\s*19\s+الإيرادات", "السياسات المحاسبية"),
    SectionDef("policy_admin_expenses", "2-20 المصاريف الإدارية والعمومية", r"2\s*-\s*20\s+.*المصاريف", "السياسات المحاسبية"),
    SectionDef("policy_financial_instruments", "2-21 الأدوات المالية", r"2\s*-\s*21\s+الأدوات", "السياسات المحاسبية"),
    SectionDef("policy_ecl", "2-22 انخفاض قيمة الأدوات المالية", r"2\s*-\s*22\s+انخفاض", "السياسات المحاسبية"),
    SectionDef("policy_ecl_measurement", "2-23 قياس الخسائر الائتمانية المتوقعة", r"2\s*-\s*23\s+قياس", "السياسات المحاسبية"),
    SectionDef("policy_ecl_presentation", "2-24 عرض مخصص الخسائر الائتمانية المتوقعة", r"2\s*-\s*24\s+عرض", "السياسات المحاسبية"),
    SectionDef("policy_writeoff", "2-25 الشطب", r"2\s*-\s*25\s+الشطب", "السياسات المحاسبية"),
    SectionDef("policy_financial_liabilities", "2-26 المطلوبات المالية", r"2\s*-\s*26\s+المطلوبات", "السياسات المحاسبية"),
]

NOTE_SECTIONS: List[SectionDef] = [
    SectionDef("note_cash", "(4) النقد وما في حكمه", r"\(\s*4\s*\)\s*النقد", "إيضاحات مالية"),
    SectionDef("note_receivables", "(5) الذمم المدينة", r"\(\s*5\s*\)\s*الذمم", "إيضاحات مالية"),
    SectionDef("note_inventory", "(6) المخزون", r"\(\s*6\s*\)\s*المخزون", "إيضاحات مالية"),
    SectionDef("note_other_debit", "(7) الأرصدة المدينة الأخرى", r"\(\s*7\s*\)\s*الأرصدة المدينة", "إيضاحات مالية"),
    SectionDef("note_ppe", "(8) الممتلكات والمعدات", r"\(\s*8\s*\)\s*الممتلكات", "إيضاحات مالية"),
    SectionDef("note_intangibles", "(9) الموجودات غير الملموسة", r"\(?\(\s*9\s*\)?\s*الموجودات غير الملموسة", "إيضاحات مالية"),
    SectionDef("note_payables", "(10) الذمم الدائنة", r"\(\s*10\s*\)\s*الذمم الدائنة", "إيضاحات مالية"),
    SectionDef("note_bank_overdraft", "(11) بنك دائن", r"\(\s*11\s*\)\s*بنك", "إيضاحات مالية"),
    SectionDef("note_accruals", "(12) المصاريف المستحقة", r"\(\s*12\s*\)\s*المصاريف", "إيضاحات مالية"),
    SectionDef("note_related_party_payable", "(13) مطلوب لأطراف ذات علاقة", r"\(\s*13\s*\)\s*مطلوب", "إيضاحات مالية"),
    SectionDef("note_postdated_checks", "(14) الشيكات الآجلة", r"\(\s*14\s*\)\s*الشيكات", "إيضاحات مالية"),
    SectionDef("note_income_tax", "(15) مخصص ضريبة الدخل والمساهمة الوطنية", r"\(\s*15\s*\)\s*مخصص", "إيضاحات مالية"),
    SectionDef("note_other_credit", "(16) الأرصدة الدائنة الأخرى", r"\(\s*16\s*\)\s*الأرصدة الدائنة", "إيضاحات مالية"),
    SectionDef("note_shareholder_loan", "(17) قرض مساهم", r"\(\s*17\s*\)\s*قرض", "إيضاحات مالية"),
    SectionDef("note_capital", "(18) رأس المال", r"\(\s*18\s*\)\s*ر[أا]س", "إيضاحات مالية"),
    SectionDef("note_statutory_reserve", "(19) الاحتياطي الإجباري", r"\(\s*19\s*\)\s*الاحتياطي", "إيضاحات مالية"),
    SectionDef("note_sales", "(20) صافي المبيعات", r"\(\s*20\s*\)\s*صافي", "إيضاحات مالية"),
    SectionDef("note_admin_expenses", "(21) المصاريف الإدارية والعمومية", r"\(\s*21\s*\)\s*المصاريف", "إيضاحات مالية"),
    SectionDef("note_contingencies", "(22) التزامات محتملة", r"\(\s*22\s*\)\s*التزامات", "إيضاحات مالية"),
]

ALL_SECTIONS: List[SectionDef] = POLICY_SECTIONS + NOTE_SECTIONS


def _clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _element_text(element) -> str:
    try:
        return _clean_text("".join(element.xpath('.//w:t/text()')))
    except Exception:
        return ""


def delete_element(element) -> None:
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _paragraphs_in_doc(doc: DocxDocument) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    for row2 in nested.rows:
                        for cell2 in row2.cells:
                            for p in cell2.paragraphs:
                                yield p
    for section in doc.sections:
        for container in (section.header, section.footer):
            for p in container.paragraphs:
                yield p
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            yield p


def _tables_in_doc(doc: DocxDocument) -> Iterable[Table]:
    for table in doc.tables:
        yield table
    for section in doc.sections:
        for container in (section.header, section.footer):
            for table in container.tables:
                yield table




def _body_children(doc: DocxDocument) -> List[object]:
    return list(doc.element.body.iterchildren())


def _find_section_ranges(doc: DocxDocument) -> Dict[str, Tuple[int, int]]:
    """Find body-element ranges for the known template sections."""
    children = _body_children(doc)
    start_indices: Dict[str, int] = {}
    for i, el in enumerate(children):
        text = _element_text(el)
        if not text:
            continue
        for sec in ALL_SECTIONS:
            if sec.key not in start_indices and re.search(sec.pattern, text):
                start_indices[sec.key] = i

    sorted_starts: List[Tuple[int, str]] = sorted((idx, key) for key, idx in start_indices.items())
    ranges: Dict[str, Tuple[int, int]] = {}
    for pos, (start, key) in enumerate(sorted_starts):
        if pos + 1 < len(sorted_starts):
            end = sorted_starts[pos + 1][0]
        else:
            end = len(children) - 1  # keep final sectPr
        ranges[key] = (start, end)
    return ranges

def remove_empty_page_break_paragraphs(doc: DocxDocument) -> None:
    """
    Remove empty paragraphs that contain only page breaks.
    This helps remove accidental blank pages caused by the template.
    """

    body = doc.element.body
    children = list(body.iterchildren())

    for el in children:
        if el.tag != qn("w:p"):
            continue

        text = _element_text(el)
        xml = el.xml

        has_page_break = (
            'w:type="page"' in xml
            or '<w:br w:type="page"' in xml
            or "<w:br/>" in xml
        )

        has_picture = "w:drawing" in xml or "w:pict" in xml

        if not text and has_page_break and not has_picture:
            delete_element(el)


def extract_policy_section_texts(template_path: Path = TEMPLATE_PATH) -> Dict[str, str]:
    """Return the default text of each accounting policy section from the template.

    The app uses these strings as the default content in editable text boxes.
    If an employee leaves a box unchanged, the generator keeps the original Word
    formatting. If they change it, that section is rebuilt from the edited text.
    """
    doc = Document(str(template_path))
    children = _body_children(doc)
    ranges = _find_section_ranges(doc)
    result: Dict[str, str] = {}
    for sec in POLICY_SECTIONS:
        if sec.key not in ranges:
            result[sec.key] = ""
            continue
        start, end = ranges[sec.key]
        pieces: List[str] = []
        for el in children[start:end]:
            text = _element_text(el)
            if text:
                pieces.append(text)
        result[sec.key] = "\n\n".join(pieces).strip()
    return result


def _looks_like_continuation(text: str, previous_item: str, blank_since_previous: bool) -> bool:
    """Return True when a Word line is probably a wrapped continuation.

    Some Arabic bullet points in the template are split across two Word paragraphs
    because of layout. We group those continuation lines into the same editable
    text box so the employee edits one complete point, not broken visual lines.
    """
    t = (text or "").strip()
    prev = (previous_item or "").strip()
    if not t or not prev:
        return False
    # Never merge a body paragraph into a short accounting-policy heading.
    if re.match(r"^2\s*-\s*\d+", prev) and len(prev) < 140:
        return False
    if blank_since_previous:
        return False
    if t.startswith("-") or t.startswith("•"):
        return False
    if prev.startswith("-") or prev.startswith("•"):
        return True
    # If the previous visual line does not end like a complete sentence, join it.
    if not re.search(r"[.،؛:)]$", prev):
        return True
    return False


def _split_heading_and_body_if_needed(text: str, sec: SectionDef) -> List[str]:
    """Split cases where Word stores the section title and first paragraph together.

    Example from this template:
    "2-4 العمليات والأرصدةيتـم ترجمة ..."
    should become two editable boxes:
    ["2-4 العمليات والأرصدة", "يتـم ترجمة ..."]
    """
    value = (text or "").strip()
    if not value:
        return []

    # Common clean case: exact display title at the start.
    if value.startswith(sec.title) and len(value) > len(sec.title):
        rest = value[len(sec.title):].strip()
        return [sec.title, rest] if rest else [sec.title]

    # Arabic/spacing variants: split after the section number and heading words
    # when the next word looks like the start of the paragraph body.
    m = re.match(
        r"^(2\s*-\s*\d+\s*[^.،؛:\n]{2,90}?)(?=(?:يتم|يتـم|تدرج|تظهر|تشمل|الموجودات|عند|تعترف|تقاس|تعتبر|إن|ان|تلغي|تحسب|تكاليف|تمثل|تتحقق|لغرض|عندما|لا\s+يتم|-\s))(.+)$",
        value,
    )
    if m:
        heading = m.group(1).strip()
        body = m.group(2).strip()
        if body:
            return [heading, body]
    return [value]


def extract_policy_section_items(template_path: Path = TEMPLATE_PATH) -> Dict[str, List[str]]:
    """Return each accounting policy as separate editable items.

    The first item is the policy title, e.g. "2-1 أسس الإعداد". Each following
    item is a separate point/paragraph. This feeds the Streamlit UI so employees
    can edit every point in its own text box.
    """
    doc = Document(str(template_path))
    children = _body_children(doc)
    ranges = _find_section_ranges(doc)
    result: Dict[str, List[str]] = {}

    for sec in POLICY_SECTIONS:
        if sec.key not in ranges:
            result[sec.key] = [sec.title]
            continue
        start, end = ranges[sec.key]
        items: List[str] = []
        blank_since_previous = False

        for el in children[start:end]:
            raw_text = _element_text(el)
            if not raw_text:
                blank_since_previous = True
                continue

            parts = _split_heading_and_body_if_needed(raw_text, sec) if not items else [raw_text]
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                if items and _looks_like_continuation(part, items[-1], blank_since_previous):
                    items[-1] = items[-1].rstrip() + "\n" + part
                else:
                    items.append(part)
                blank_since_previous = False

        result[sec.key] = items or [sec.title]
    return result


def _make_arabic_paragraph_element(text: str, bold: bool = False):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)
    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p


def apply_policy_text_edits(doc: DocxDocument, policy_text_edits: Dict[str, str]) -> None:
    """Replace edited accounting policy sections with employee text.

    Only sections included in policy_text_edits are rebuilt. Unchanged sections
    keep their original Word layout. This is intentional because some policies
    contain tables; those are only rebuilt when the employee actually edits the
    text box for that section.
    """
    if not policy_text_edits:
        return
    body = doc.element.body
    children = _body_children(doc)
    ranges = _find_section_ranges(doc)

    # Work from the bottom of the document upward so earlier indexes remain valid.
    replacements: List[Tuple[int, int, str]] = []
    for key, text in policy_text_edits.items():
        if key not in ranges:
            continue
        start, end = ranges[key]
        replacements.append((start, end, text or ""))
    replacements.sort(reverse=True, key=lambda x: x[0])

    for start, end, text in replacements:
        paragraphs = [line.strip() for line in str(text).splitlines()]
        paragraphs = [line for line in paragraphs if line]
        new_elements = []
        if not paragraphs:
            new_elements = [_make_arabic_paragraph_element("")]
        else:
            for i, line in enumerate(paragraphs):
                new_elements.append(_make_arabic_paragraph_element(line, bold=(i == 0)))

        # Delete old section elements.
        for el in children[start:end]:
            delete_element(el)
        # Insert rebuilt section at the same position.
        for offset, new_el in enumerate(new_elements):
            body.insert(start + offset, new_el)


def remove_red_and_highlight_everywhere(doc: DocxDocument) -> None:
    """Clean final document markup: remove red font and yellow highlights."""
    for p in _paragraphs_in_doc(doc):
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.highlight_color = None

def replace_in_paragraph(paragraph: Paragraph, replacements: Dict[str, str], clear_replaced_format: bool = True) -> None:
    if not paragraph.runs:
        return

    # First: safe run-level replacement.
    for run in paragraph.runs:
        if not run.text:
            continue
        new_text = run.text
        changed = False
        for old, new in replacements.items():
            if old and old in new_text:
                new_text = new_text.replace(old, str(new))
                changed = True
        if changed:
            run.text = new_text
            if clear_replaced_format:
                run.font.color.rgb = None
                run.font.highlight_color = None

    # Second: paragraph-level fallback for text split across several runs.
    full_text = paragraph.text
    new_full = full_text
    for old, new in replacements.items():
        if old and old in new_full:
            new_full = new_full.replace(old, str(new))
    if new_full != full_text:
        # Keep paragraph style/alignment. Preserve formatting of first run only.
        first_run = paragraph.runs[0]
        first_run.text = new_full
        if clear_replaced_format:
            first_run.font.color.rgb = None
            first_run.font.highlight_color = None
        for r in paragraph.runs[1:]:
            r.text = ""


def apply_replacements(doc: DocxDocument, replacements: Dict[str, str], clear_replaced_format: bool = True) -> None:
    # Longest first prevents partial replacement from eating a longer phrase.
    ordered = dict(sorted(replacements.items(), key=lambda kv: len(kv[0]), reverse=True))
    for p in _paragraphs_in_doc(doc):
        replace_in_paragraph(p, ordered, clear_replaced_format)


def set_cell_text(cell: _Cell, text: str) -> None:
    text = "" if text is None else str(text)
    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    # Remove extra paragraphs inside cell, but keep XML valid.
    for extra in cell.paragraphs[1:]:
        delete_element(extra._element)

    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.highlight_color = None

def make_ltr(text: str) -> str:
    text = "" if text is None else str(text)
    if not text.strip():
        return ""
    return f"\u200e{text}\u200e"


def to_float(value) -> float:
    try:
        cleaned = str(value or "").replace(",", "").replace("%", "").strip()
        if not cleaned:
            return 0.0
        return float(cleaned)
    except ValueError:
        return 0.0


def format_percent_for_doc(value: float) -> str:
    if not value:
        return ""
    if value == int(value):
        return f"{int(value)}%"
    return f"{value:.2f}%"


def set_table_rtl(table: Table) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)

    if tbl_pr.find(qn("w:bidiVisual")) is None:
        tbl_pr.append(OxmlElement("w:bidiVisual"))


def set_cell_width(cell: _Cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)

    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_no_wrap(cell: _Cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    if tc_pr.find(qn("w:noWrap")) is None:
        tc_pr.append(OxmlElement("w:noWrap"))


def set_cell_bottom_border(cell: _Cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    bottom = tc_borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        tc_borders.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), "000000")


def set_cell_top_border(cell: _Cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    top = tc_borders.find(qn("w:top"))
    if top is None:
        top = OxmlElement("w:top")
        tc_borders.append(top)

    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), "000000")


def write_clean_cell(
    cell: _Cell,
    text: str,
    width_dxa: int,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    no_wrap: bool = True,
) -> None:
    set_cell_text(cell, text)
    set_cell_width(cell, width_dxa)

    if no_wrap:
        set_cell_no_wrap(cell)

    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1

        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)


def fill_auditor_table(doc: DocxDocument, form: Dict[str, str]) -> None:
    """Fill the auditor block without changing the partners-table header."""
    city_country = f"{form.get('city','')} - {form.get('country','')}".strip(" -")
    if not doc.tables:
        return
    # In this template the auditor block is the first table.
    table = doc.tables[0]
    if len(table.rows) > 0 and len(table.rows[0].cells) > 2:
        set_cell_text(table.rows[0].cells[2], form.get("audit_office", ""))
    if len(table.rows) > 3 and len(table.rows[3].cells) > 2:
        set_cell_text(table.rows[3].cells[2], form.get("audit_partner", ""))
    if len(table.rows) > 4 and len(table.rows[4].cells) > 2:
        lic = form.get("audit_license", "")
        set_cell_text(table.rows[4].cells[2], f"اجازة رقم ({lic})" if lic else "اجازة رقم ()")
    if len(table.rows) > 5 and len(table.rows[5].cells) > 2:
        set_cell_text(table.rows[5].cells[2], city_country)
    if len(table.rows) > 6 and len(table.rows[6].cells) > 2:
        set_cell_text(table.rows[6].cells[2], f"التاريخ: {form.get('audit_date','')}")


def fill_partners_table(doc: DocxDocument, partners: List[Dict[str, str]]) -> None:
    if not partners:
        return

    target = None

    for table in doc.tables:
        txt = _clean_text(" ".join(cell.text for row in table.rows for cell in row.cells))

        if "اسم الشريك" in txt and "عدد الحصص" in txt and "نسبة المساهمة" in txt:
            target = table
            break

    if target is None:
        return

    total_shares = 0.0
    total_value = 0.0
    total_percentage = 0.0

    for partner in partners:
        total_shares += to_float(partner.get("shares", ""))
        total_value += to_float(partner.get("value", ""))
        total_percentage += to_float(partner.get("percentage", ""))

    # Create a fresh clean table at the end first
    new_table = doc.add_table(rows=len(partners) + 2, cols=4)
    new_table.autofit = False
    set_table_rtl(new_table)

    # Widths are in twentieths of a point.
    # Because the table is RTL:
    # cell 0 appears on the right, cell 3 appears on the left.
    widths = {
        "name": 3600,
        "shares": 1800,
        "value": 2000,
        "percentage": 1800,
    }

    # Header row
    header = new_table.rows[0]
    write_clean_cell(header.cells[0], "اسم الشريك", widths["name"], bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    write_clean_cell(header.cells[1], "عدد الحصص", widths["shares"], bold=True)
    write_clean_cell(header.cells[2], "قيمة الحصص\nدينار أردني", widths["value"], bold=True)
    write_clean_cell(header.cells[3], "نسبة المساهمة", widths["percentage"], bold=True)

    for cell in header.cells:
        set_cell_bottom_border(cell)

    # Partner rows
    for i, partner in enumerate(partners, start=1):
        row = new_table.rows[i]

        write_clean_cell(
            row.cells[0],
            partner.get("name", ""),
            widths["name"],
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            no_wrap=False,
        )

        write_clean_cell(row.cells[1], partner.get("shares", ""), widths["shares"])
        write_clean_cell(row.cells[2], partner.get("value", ""), widths["value"])
        write_clean_cell(row.cells[3], make_ltr(partner.get("percentage", "")), widths["percentage"])

    # Total row
    total_row = new_table.rows[len(partners) + 1]

    write_clean_cell(total_row.cells[0], "المجموع", widths["name"], bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    write_clean_cell(total_row.cells[1], f"{total_shares:,.0f}" if total_shares else "", widths["shares"], bold=True)
    write_clean_cell(total_row.cells[2], f"{total_value:,.0f}" if total_value else "", widths["value"], bold=True)
    write_clean_cell(total_row.cells[3], make_ltr(format_percent_for_doc(total_percentage)), widths["percentage"], bold=True)

    for cell in total_row.cells:
        set_cell_top_border(cell)
        set_cell_bottom_border(cell)

    # Move the new clean table to the old table position, then remove the old broken table
    old_tbl = target._tbl
    parent = old_tbl.getparent()
    old_index = parent.index(old_tbl)

    new_tbl = new_table._tbl
    new_tbl.getparent().remove(new_tbl)

    parent.insert(old_index, new_tbl)
    parent.remove(old_tbl)


def extract_table_data(template_path: Path = TEMPLATE_PATH) -> List[Dict[str, object]]:
    doc = Document(str(template_path))
    result = []
    for idx, table in enumerate(doc.tables):
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        # Short label based on first non-empty cells.
        preview_cells = [_clean_text(c) for r in data[:4] for c in r if _clean_text(c)]
        label = " | ".join(preview_cells[:3])
        if not label:
            label = f"Table {idx + 1}"
        result.append({"index": idx, "label": label[:120], "data": data})
    return result


def remove_table_row(table, row_index: int) -> None:
    row = table.rows[row_index]
    tbl = table._tbl
    tbl.remove(row._tr)


def remove_table_column(table, col_index: int) -> None:
    for row in table.rows:
        cell = row.cells[col_index]
        tc = cell._tc
        tc.getparent().remove(tc)

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is not None:
        grid_cols = list(tbl_grid.gridCol_lst)
        if col_index < len(grid_cols):
            tbl_grid.remove(grid_cols[col_index])


def resize_word_table(table, target_rows: int, target_cols: int) -> None:
    target_rows = max(1, int(target_rows))
    target_cols = max(1, int(target_cols))

    while len(table.rows) < target_rows:
        table.add_row()

    while len(table.rows) > target_rows:
        remove_table_row(table, len(table.rows) - 1)

    current_cols = len(table.rows[0].cells) if table.rows else 0

    while current_cols < target_cols:
        table.add_column(914400)  # about 1 inch in EMUs
        current_cols += 1

    while current_cols > target_cols:
        remove_table_column(table, current_cols - 1)
        current_cols -= 1


def apply_table_data(doc, table_updates):
    if not table_updates:
        return

    for table_index, data in table_updates.items():
        table_index = int(table_index)

        if table_index < 0 or table_index >= len(doc.tables):
            continue

        if not data:
            continue

        table = doc.tables[table_index]

        # Add rows only if needed. Do NOT add/remove columns.
        while len(table.rows) < len(data):
            table.add_row()

        # Fill replacement rows
        for row_index, row_data in enumerate(data):
            if row_index >= len(table.rows):
                continue

            row = table.rows[row_index]

            # Clear the row first
            for cell in row.cells:
                set_cell_text(cell, "")

            # Fill only existing Word cells
            max_cols = min(len(row.cells), len(row_data))
            for col_index in range(max_cols):
                set_cell_text(row.cells[col_index], row_data[col_index])

        # Clear leftover old rows below the replacement data
        for row_index in range(len(data), len(table.rows)):
            row = table.rows[row_index]
            for cell in row.cells:
                set_cell_text(cell, "")




def _cell_matches_label(cell_text: str, label: str) -> bool:
    """Return True when a Word table first-cell label matches the target line item.

    Conservative matching is important. For example, the row label "المبيعات"
    must NOT match the target "صافي المبيعات". So we allow exact matching,
    and allow the target label to appear inside a longer Word row label such as
    "ينزل: خصم المبيعات". We do not allow the reverse direction.
    """
    cell_norm = _clean_text(cell_text)
    label_norm = _clean_text(label)
    if not cell_norm or not label_norm:
        return False
    return cell_norm == label_norm or label_norm in cell_norm


def _find_row_by_label(table, label: str) -> Optional[int]:
    """Find a table row by the visible first-column Arabic label."""
    if not label:
        return None
    for idx, row in enumerate(table.rows):
        if not row.cells:
            continue
        if _cell_matches_label(row.cells[0].text, label):
            return idx
    return None


def apply_financial_note_values(doc: DocxDocument, financial_note_values: Dict[str, Dict[str, str]]) -> None:
    """Apply extracted/edited numeric note values to the Word note tables.

    V7 fix: V6 relied only on hard-coded table/row indexes. That can update the
    wrong row if the template changes slightly, or if a row is shifted. This
    version first tries to find the row by the Arabic line-item label, then falls
    back to the old row number.

    Keys are strings in the format "table_index:row_index". Values contain:
    {"current": "...", "previous": "...", "label": "..."}.
    """
    if not financial_note_values:
        return
    for key, values in financial_note_values.items():
        try:
            table_idx_s, row_idx_s = str(key).split(":", 1)
            table_idx = int(table_idx_s)
            fallback_row_idx = int(row_idx_s)
        except Exception:
            continue
        if table_idx < 0 or table_idx >= len(doc.tables):
            continue
        table = doc.tables[table_idx]

        label = str(values.get("label", "") or "")
        found_idx = _find_row_by_label(table, label)
        row_idx = found_idx if found_idx is not None else fallback_row_idx

        if row_idx < 0 or row_idx >= len(table.rows):
            continue
        row = table.rows[row_idx]
        if len(row.cells) > 1:
            set_cell_text(row.cells[1], values.get("current", ""))
        if len(row.cells) > 3:
            set_cell_text(row.cells[3], values.get("previous", ""))



def apply_cell_values(doc: DocxDocument, cell_values: Dict[str, str]) -> None:
    """Apply exact table-cell updates, used by V10 movement tables.

    Keys are "table_index:row_index:col_index".
    This is intentionally separate from financial note rows because movement
    tables have many columns and do not follow the simple 2024/2023 row format.
    """
    if not cell_values:
        return
    for key, value in cell_values.items():
        try:
            table_idx_s, row_idx_s, col_idx_s = str(key).split(":", 2)
            table_idx = int(table_idx_s)
            row_idx = int(row_idx_s)
            col_idx = int(col_idx_s)
        except Exception:
            continue
        if table_idx < 0 or table_idx >= len(doc.tables):
            continue
        table = doc.tables[table_idx]
        if row_idx < 0 or row_idx >= len(table.rows):
            continue
        row = table.rows[row_idx]
        if col_idx < 0 or col_idx >= len(row.cells):
            continue
        set_cell_text(row.cells[col_idx], value)


def remove_sections(doc: DocxDocument, section_keys_to_remove: Iterable[str]) -> None:
    section_keys_to_remove = set(section_keys_to_remove)
    if not section_keys_to_remove:
        return

    body = doc.element.body
    children = list(body.iterchildren())
    start_indices: Dict[str, int] = {}

    for i, el in enumerate(children):
        text = _element_text(el)
        if not text:
            continue
        for sec in ALL_SECTIONS:
            if sec.key not in start_indices and re.search(sec.pattern, text):
                start_indices[sec.key] = i

    # Create a sorted list of all known starts so each section ends at the next known start.
    sorted_starts: List[Tuple[int, str]] = sorted((idx, key) for key, idx in start_indices.items())
    ranges: List[Tuple[int, int]] = []
    for pos, (start, key) in enumerate(sorted_starts):
        if key not in section_keys_to_remove:
            continue
        if pos + 1 < len(sorted_starts):
            end = sorted_starts[pos + 1][0]
        else:
            end = len(children) - 1  # keep final sectPr
        ranges.append((start, end))

    # Merge overlapping ranges and delete from back to front.
    ranges.sort()
    merged: List[Tuple[int, int]] = []
    for start, end in ranges:
        if not merged or start > merged[-1][1]:
            merged.append((start, end))
        else:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))

    for start, end in reversed(merged):
        for el in children[start:end]:
            delete_element(el)


def append_custom_sections(doc: DocxDocument, sections: List[Dict[str, str]]) -> None:
    for sec in sections:
        title = _clean_text(sec.get("title", ""))
        body = sec.get("body", "") or ""
        if not title and not body.strip():
            continue
        doc.add_page_break()
        if title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(14)
        for para in body.splitlines():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(para)


def build_replacements(form: Dict[str, str]) -> Dict[str, str]:
    city_country = f"{form.get('city','')} - {form.get('country','')}".strip(" -")
    threshold_raw = form.get("important_partner_threshold", "5")

    try:
        threshold_num = float(threshold_raw)
        if threshold_num == int(threshold_num):
            threshold_text = str(int(threshold_num))
        else:
            threshold_text = str(threshold_num).rstrip("0").rstrip(".")
    except (TypeError, ValueError):
        threshold_text = str(threshold_raw or "5").replace("%", "").strip() or "5"
    replacements = {
        'اسم الشركة "من اسم الملف الاساسي"': form.get("company_name", ""),
        'اسم الشركة من اسم الملف الاساسي': form.get("company_name", ""),
        '(نوع الشركة من الملف الاساسي)': f"({form.get('company_type','')})" if form.get("company_type") else "",
        'نوع الشركة من الملف الاساسي': form.get("company_type", ""),
        'المدينة  - البلد "من الملف الاساسي"': city_country,
        'المدينة - البلد من الملف الاساسي': city_country,
        'المدينة  - البلد من الملف الاساسي': city_country,
        'السنة المالية من الملف الاساسي': form.get("financial_year", ""),
        
        'اسم مكتب التدقيق': form.get("audit_office", ""),
        'اجازة رقم (.........)': f"اجازة رقم ({form.get('audit_license','')})" if form.get("audit_license") else "اجازة رقم ()",
        'المدينة – البلد': city_country,
        'التاريخ :': f"التاريخ : {form.get('audit_date','')}",
        'تأسست اسم الشركة': f"تأسست {form.get('company_name','')}",
        '(نوع الشركة)': f"({form.get('company_type','')})",
        'في المدينة': f"في {form.get('city','')}",
        '(رقم التسجيل من ملف الشركة)': f"({form.get('registration_number','')})",
        'تاريخ التسجيل من ملف الشركة': form.get("registration_date", ""),
        'رأس المال من ملف الشركة': form.get("capital", ""),
        'العملة من ملف الشركة': form.get("currency", ""),
        'صندوق بريد  ... من ملف الشركة.....اسم المدينة – الرمز البريدي من ملف الشركة - الدولة': f"صندوق بريد {form.get('po_box','')} {form.get('city','')} – {form.get('postal_code','')} - {form.get('country','')}",
        'بتاريخ ..........': f"بتاريخ {form.get('approval_date','')}",
        'ان اهم الشركاء في الشركة والذين تزيد نسبة مساهمتهم عن 5% هم كما يلي :': f'ان اهم الشركاء في الشركة والذين تزيد نسبة مساهمتهم عن {threshold_text}% هم كما يلي :',
        'ان اهم الشركاء في الشركة والذين تزيد نسبة مساهمتهم عن 5 هم كما يلي % :': f'ان اهم الشركاء في الشركة والذين تزيد نسبة مساهمتهم عن {threshold_text}% هم كما يلي :',
        'والذين تزيد نسبة مساهمتهم عن 5% هم كما يلي': f'والذين تزيد نسبة مساهمتهم عن {threshold_text}% هم كما يلي',
        'والذين تزيد نسبة مساهمتهم عن 5 هم كما يلي %': f'والذين تزيد نسبة مساهمتهم عن {threshold_text}% هم كما يلي',
    }
    # Do NOT globally replace "اسم الشريك" because it is also the partners table header.
    return {k: v for k, v in replacements.items() if v is not None}


def insert_audit_logo_on_report_pages(doc, logo_path: str, width_inches: float = 0.70) -> None:
    """
    Replace 'شعار مكتب التدقيق' with the audit logo.
    Searches normal paragraphs, tables, headers, and footers.
    """

    logo = Path(logo_path)

    if not logo.exists():
        return

    target_text = "شعار مكتب التدقيق"
    replaced_count = 0

    for paragraph in _paragraphs_in_doc(doc):
        if paragraph.text.strip() == target_text:
            for run in paragraph.runs:
                run.text = ""

            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = paragraph.add_run()
            run.add_picture(str(logo), width=Inches(width_inches))

            replaced_count += 1

            if replaced_count >= 2:
                break


def fix_blank_pages_from_section_breaks(doc: DocxDocument) -> None:
    """
    Safer blank-page fix:
    - Converts odd/even section breaks to normal nextPage breaks.
    - Removes only paragraphs that contain a page break and no real content.
    - Does NOT remove normal empty paragraphs, because those protect Arabic table layout.
    """

    body = doc.element.body

    # 1) Convert odd/even section breaks to normal nextPage breaks.
    for sectPr in doc.element.xpath(".//w:sectPr"):
        type_el = sectPr.find(qn("w:type"))
        if type_el is not None:
            value = type_el.get(qn("w:val"))
            if value in ("oddPage", "evenPage"):
                type_el.set(qn("w:val"), "nextPage")

    # 2) Remove only empty paragraphs that contain a manual page break.
    for el in list(body.iterchildren()):
        if el.tag != qn("w:p"):
            continue

        text = _element_text(el)
        xml = el.xml

        has_page_break = (
            'w:type="page"' in xml
            or '<w:br w:type="page"' in xml
        )

        has_picture = "w:drawing" in xml or "w:pict" in xml
        has_section_properties = "w:sectPr" in xml

        if not text and has_page_break and not has_picture and not has_section_properties:
            delete_element(el)


def remove_r4_financial_statement_title(doc: DocxDocument) -> None:
    """
    Remove all unwanted R4 / 4R report title text.
    Works in normal paragraphs, tables, headers, and footers.
    """

    patterns = [
        r"R4\s*[اإ]يضاحات\s+القوائم\s+المالية",
        r"4R\s*[اإ]يضاحات\s+القوائم\s+المالية",
        r"[اإ]يضاحات\s+القوائم\s+المالية\s*R4",
        r"[اإ]يضاحات\s+القوائم\s+المالية\s*4R",

        # Remove من التقرير R4 / من التقرير 4R
        r"من\s+التقرير\s*R4",
        r"من\s+التقرير\s*4R",
        r"R4\s*من\s+التقرير",
        r"4R\s*من\s+التقرير",
    ]

    for paragraph in list(_paragraphs_in_doc(doc)):
        original_text = paragraph.text or ""
        new_text = original_text

        for pattern in patterns:
            new_text = re.sub(pattern, "", new_text).strip()

        if new_text != original_text:
            if not new_text:
                delete_element(paragraph._element)
            else:
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for run in paragraph.runs[1:]:
                        run.text = ""
                else:
                    paragraph.add_run(new_text)


def generate_document(
    form: Dict[str, str],
    partners: List[Dict[str, str]],
    included_section_keys: Iterable[str],
    custom_sections: Optional[List[Dict[str, str]]] = None,
    table_updates: Optional[Dict[int, List[List[str]]]] = None,
    financial_note_values: Optional[Dict[str, Dict[str, str]]] = None,
    cell_values: Optional[Dict[str, str]] = None,
    policy_text_edits: Optional[Dict[str, str]] = None,
    clear_replaced_format: bool = True,
) -> bytes:
    doc = Document(str(TEMPLATE_PATH))

    replacements = build_replacements(form)
    apply_replacements(doc, replacements, clear_replaced_format)
    update_cover_page(doc, form)

    fill_auditor_table(doc, form)
    fill_partners_table(doc, partners)

    if policy_text_edits:
        apply_policy_text_edits(doc, policy_text_edits)

    # IMPORTANT:
    # Apply manual table replacements BEFORE removing sections,
    # because table indexes come from the original Word template.
    if table_updates:
        apply_table_data(doc, table_updates)

    included = set(included_section_keys)
    all_keys = {s.key for s in ALL_SECTIONS}
    remove_keys = all_keys - included
    remove_sections(doc, remove_keys)

    if financial_note_values:
        apply_financial_note_values(doc, financial_note_values)

    if cell_values:
        apply_cell_values(doc, cell_values)

    append_custom_sections(doc, custom_sections or [])

    insert_audit_logo_on_report_pages(
        doc,
        logo_path="assets/logo.png",
        width_inches=0.70
    )

    fix_blank_pages_from_section_breaks(doc)

    remove_r4_financial_statement_title(doc)

    if clear_replaced_format:
        remove_red_and_highlight_everywhere(doc)

    update_page_two_and_index(doc, form)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
