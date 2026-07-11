from __future__ import annotations

from typing import Iterable, List, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


MAIN_INDEX_ROWS = [
    ("تقرير مدقق الحسابات المستقل", ""),
    ("قائمة المركز المالي", "1"),
    ("قائمة الدخل الشامل", "2"),
    ("قائمة التغيـّـرات فـي حقوق الملكيّـة", "3"),
    ("قائمة التدفقـات النقديــة", "4"),
]


NOTE_INDEX_ROWS = [
    ("note_cash", "(4) النقد وما في حكمه"),
    ("note_receivables", "(5) الذمم المدينة"),
    ("note_inventory", "(6) المخزون"),
    ("note_other_debit", "(7) الأرصدة المدينة الأخرى"),
    ("note_ppe", "(8) الممتلكات والمعدات"),
    ("note_intangibles", "(9) الموجودات غير الملموسة"),
    ("note_payables", "(10) الذمم الدائنة"),
    ("note_bank_overdraft", "(11) بنك دائن"),
    ("note_accruals", "(12) المصاريف المستحقة"),
    ("note_related_party_payable", "(13) مطلوب لأطراف ذات علاقة"),
    ("note_postdated_checks", "(14) الشيكات الآجلة"),
    ("note_income_tax", "(15) مخصص ضريبة الدخل والمساهمة الوطنية"),
    ("note_other_credit", "(16) الأرصدة الدائنة الأخرى"),
    ("note_shareholder_loan", "(17) قرض مساهم"),
    ("note_capital", "(18) رأس المال"),
    ("note_statutory_reserve", "(19) الاحتياطي الإجباري"),
    ("note_sales", "(20) صافي المبيعات"),
    ("note_admin_expenses", "(21) المصاريف الإدارية والعمومية"),
    ("note_contingencies", "(22) التزامات محتملة"),
]


def set_paragraph_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()

    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)

    bidi.set(qn("w:val"), "1")


def clean_text(value) -> str:
    return " ".join(str(value or "").split())


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    text = str(text or "")

    if not cell.paragraphs:
        paragraph = cell.add_paragraph()
    else:
        paragraph = cell.paragraphs[0]

    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)

    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_rtl(paragraph)

    for run in paragraph.runs:
        run.bold = bold


def remove_table_row(table, row_index: int) -> None:
    row = table.rows[row_index]
    table._tbl.remove(row._tr)


def find_index_table(doc):
    """
    Finds the index table on page 2.
    It looks for the table containing صفحة and قائمة المركز المالي.
    """
    for table in doc.tables:
        text = clean_text(" ".join(cell.text for row in table.rows for cell in row.cells))

        if "صفح" in text and "قائمة المركز المالي" in text:
            return table

    return None


def build_index_rows(included_section_keys: Iterable[str]) -> List[Tuple[str, str]]:
    included = set(included_section_keys)

    rows: List[Tuple[str, str]] = []
    rows.extend(MAIN_INDEX_ROWS)

    page_number = 5

    for key, title in NOTE_INDEX_ROWS:
        if key in included:
            rows.append((title, str(page_number)))
            page_number += 1

    return rows


def update_index_table(doc, included_section_keys: Iterable[str]) -> None:
    """
    Updates the index/table of contents based on selected financial note sections.
    """
    table = find_index_table(doc)

    if table is None:
        return

    rows = build_index_rows(included_section_keys)

    needed_rows = len(rows) + 1  # +1 for header

    while len(table.rows) < needed_rows:
        table.add_row()

    while len(table.rows) > needed_rows:
        remove_table_row(table, len(table.rows) - 1)

    # Header row
    header = table.rows[0]
    if len(header.cells) >= 2:
        set_cell_text(header.cells[0], "", bold=True)
        set_cell_text(header.cells[1], "صفحـــة", bold=True)

    # Data rows
    for i, (title, page) in enumerate(rows, start=1):
        row = table.rows[i]

        if len(row.cells) >= 2:
            set_cell_text(row.cells[0], title)
            set_cell_text(row.cells[1], page)
