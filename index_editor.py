from __future__ import annotations

from copy import deepcopy
import re
from typing import Dict, Optional, Sequence

from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


_ARABIC_DIACRITICS = re.compile(
    r"[\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06ED]"
)


def _match_text(value: object) -> str:
    """Normalize Arabic text only for matching template labels."""
    text = str(value or "")
    text = text.replace("ـ", "")
    text = _ARABIC_DIACRITICS.sub("", text)
    text = (
        text.replace("أ", "ا")
        .replace("إ", "ا")
        .replace("آ", "ا")
        .replace("ى", "ي")
    )
    return re.sub(r"\s+", " ", text).strip()


def _set_rtl(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()

    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)

    bidi.set(qn("w:val"), "1")


def _replace_keep_style(paragraph: Paragraph, new_text: str) -> None:
    """Replace visible paragraph text while preserving the first run style."""
    new_text = str(new_text or "").strip()

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text

    for run in paragraph.runs[1:]:
        run.text = ""


def _find_first_index_row_position(doc: DocxDocument) -> Optional[int]:
    """
    The page-2 index rows contain a tab between the Arabic title and page number.
    """
    for index, paragraph in enumerate(doc.paragraphs[:100]):
        text = _match_text(paragraph.text)

        if "\t" in paragraph.text and "قائمة المركز المالي" in text:
            return index

    return None


def _update_page_two_heading(doc: DocxDocument, form: Dict[str, str]) -> None:
    """
    Update the title block above the index on page 2.

    Important: this block is normal document text, not a Word header.
    """
    first_index_row = _find_first_index_row_position(doc)
    if first_index_row is None:
        return

    company_name = str(form.get("company_name", "") or "").strip()

    legal_form = str(form.get("legal_form", "") or "").strip()
    if not legal_form:
        legal_form = "(ذات مسؤولية محدودة)"

    statements_title = str(form.get("statements_title", "") or "").strip()
    if not statements_title:
        statements_title = "القوائم الماليــة المنفصلة للسنة المالية المنتهية في"

    cover_date = str(form.get("cover_date", "") or "").strip()
    if not cover_date:
        cover_date = str(form.get("financial_year", "") or "").strip()

    combined_title = statements_title
    if cover_date and _match_text(cover_date) not in _match_text(statements_title):
        combined_title = f"{statements_title} {cover_date}".strip()

    # Only inspect the small heading area immediately above the page-2 index.
    start = max(0, first_index_row - 15)

    for paragraph in doc.paragraphs[start:first_index_row]:
        text = _match_text(paragraph.text)

        if text == "شركة الاحتراف للتعليم" and company_name:
            _replace_keep_style(paragraph, company_name)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(paragraph)
            continue

        if "ذات مسؤولية محدودة" in text and len(text) < 60:
            _replace_keep_style(paragraph, legal_form)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(paragraph)
            continue

        if text.startswith(
            "القوائم المالية المنفصلة للسنة المالية المنتهية في"
        ):
            _replace_keep_style(paragraph, combined_title)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_rtl(paragraph)


def _next_bookmark_id(doc: DocxDocument) -> int:
    bookmark_ids = []

    for bookmark in doc._element.xpath(".//w:bookmarkStart"):
        raw_id = bookmark.get(qn("w:id"))

        try:
            bookmark_ids.append(int(raw_id))
        except (TypeError, ValueError):
            continue

    return max(bookmark_ids, default=0) + 1


def _remove_bookmark(doc: DocxDocument, bookmark_name: str) -> None:
    starts = doc._element.xpath(
        f'.//w:bookmarkStart[@w:name="{bookmark_name}"]'
    )
    bookmark_ids = {start.get(qn("w:id")) for start in starts}

    for start in starts:
        parent = start.getparent()
        if parent is not None:
            parent.remove(start)

    for end in doc._element.xpath(".//w:bookmarkEnd"):
        if end.get(qn("w:id")) in bookmark_ids:
            parent = end.getparent()
            if parent is not None:
                parent.remove(end)


def _add_bookmark(
    doc: DocxDocument,
    paragraph: Paragraph,
    bookmark_name: str,
) -> None:
    """Place a Word bookmark on the paragraph containing a target title."""
    _remove_bookmark(doc, bookmark_name)

    bookmark_id = str(_next_bookmark_id(doc))

    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), bookmark_id)
    bookmark_start.set(qn("w:name"), bookmark_name)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), bookmark_id)

    paragraph_xml = paragraph._p
    insert_position = 1 if paragraph_xml.pPr is not None else 0

    paragraph_xml.insert(insert_position, bookmark_start)
    paragraph_xml.append(bookmark_end)


def _append_text_run(
    paragraph: Paragraph,
    text: str,
    run_properties,
):
    run = paragraph.add_run(text)

    if run_properties is not None:
        if run._r.rPr is not None:
            run._r.remove(run._r.rPr)

        run._r.insert(0, deepcopy(run_properties))

    return run


def _append_pageref_field(
    paragraph: Paragraph,
    bookmark_name: str,
    cached_result: str,
    run_properties,
) -> None:
    """
    Insert a PAGEREF field.

    Word recalculates the displayed page number from the bookmark location.
    """
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), f"PAGEREF {bookmark_name} \\h")
    field.set(qn("w:dirty"), "true")

    field_run = OxmlElement("w:r")

    if run_properties is not None:
        field_run.append(deepcopy(run_properties))

    text = OxmlElement("w:t")
    text.text = str(cached_result)

    field_run.append(text)
    field.append(field_run)
    paragraph._p.append(field)


def _replace_index_row(
    paragraph: Paragraph,
    title: str,
    page_references: Sequence[tuple[str, str]],
) -> None:
    """
    Rebuild one page-2 index row while preserving its paragraph formatting.
    """
    run_properties = None

    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        run_properties = deepcopy(paragraph.runs[0]._r.rPr)

    # Keep paragraph properties such as tabs, alignment and spacing.
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)

    _append_text_run(paragraph, title, run_properties)

    tab_run = _append_text_run(paragraph, "", run_properties)
    tab_run.add_tab()

    for reference_index, (bookmark_name, cached_result) in enumerate(
        page_references
    ):
        if reference_index > 0:
            _append_text_run(paragraph, " - ", run_properties)

        _append_pageref_field(
            paragraph,
            bookmark_name,
            cached_result,
            run_properties,
        )


def _ensure_auditor_report_index_paragraph(
    doc: DocxDocument,
) -> Optional[Paragraph]:
    """
    Ensure page 2 contains an index row for
    'تقرير مدقق الحسابات المستقل'.

    The template currently has no row for this title, so clone the
    balance-sheet index row and insert the clone immediately before it.
    """
    existing = _find_index_paragraph(
        doc,
        "تقرير مدقق الحسابات المستقل",
    )
    if existing is not None:
        return existing

    balance_sheet_row = _find_index_paragraph(
        doc,
        "قائمة المركز المالي",
    )
    if balance_sheet_row is None:
        return None

    cloned_xml = deepcopy(balance_sheet_row._p)
    balance_sheet_row._p.addprevious(cloned_xml)

    return Paragraph(cloned_xml, balance_sheet_row._parent)


def _find_index_paragraph(
    doc: DocxDocument,
    phrase: str,
) -> Optional[Paragraph]:
    phrase = _match_text(phrase)

    for paragraph in doc.paragraphs[:100]:
        text = _match_text(paragraph.text)

        if "\t" in paragraph.text and phrase in text:
            return paragraph

    return None


def _find_exact_title_paragraph(
    doc: DocxDocument,
    phrase: str,
) -> Optional[Paragraph]:
    phrase = _match_text(phrase)

    for paragraph in doc.paragraphs[100:]:
        if _match_text(paragraph.text) == phrase:
            return paragraph

    return None


def _find_notes_start(doc: DocxDocument) -> Optional[Paragraph]:
    """
    The notes title itself is in the repeating Word header.
    Bookmark the first note body title instead: '(1) معلومات عامة'.
    """
    for paragraph in doc.paragraphs[100:]:
        text = _match_text(paragraph.text)

        if "معلومات عامة" in text and len(text) < 80:
            return paragraph

    return None


def _last_paragraph_in_table(table: Table) -> Optional[Paragraph]:
    for row in reversed(table.rows):
        for cell in reversed(row.cells):
            for nested_table in reversed(cell.tables):
                nested_paragraph = _last_paragraph_in_table(nested_table)
                if nested_paragraph is not None:
                    return nested_paragraph

            if cell.paragraphs:
                return cell.paragraphs[-1]

    return None


def _find_last_body_paragraph(doc: DocxDocument) -> Optional[Paragraph]:
    """
    Find the last paragraph in document-body order, including a final table.
    This gives the notes range its automatic ending page.
    """
    for child in reversed(list(doc.element.body.iterchildren())):
        if child.tag == qn("w:p"):
            return Paragraph(child, doc._body)

        if child.tag == qn("w:tbl"):
            table = Table(child, doc._body)
            paragraph = _last_paragraph_in_table(table)

            if paragraph is not None:
                return paragraph

    return None


def _ensure_section_break_before_paragraph(
    doc: DocxDocument,
    target_paragraph: Paragraph,
) -> None:
    """
    Split the current Word section immediately before target_paragraph.

    The original template keeps the cover, index and auditor report in one
    section. PAGEREF cannot show the auditor report as page 1 unless the
    report starts its own section.
    """
    body_children = list(doc.element.body.iterchildren())

    try:
        target_position = next(
            index
            for index, child in enumerate(body_children)
            if child is target_paragraph._p
        )
    except StopIteration:
        return

    previous_paragraph_xml = None
    for child in reversed(body_children[:target_position]):
        if child.tag == qn("w:p"):
            previous_paragraph_xml = child
            break

    if previous_paragraph_xml is None:
        return

    previous_p_pr = previous_paragraph_xml.find(qn("w:pPr"))
    if previous_p_pr is not None:
        existing_section = previous_p_pr.find(qn("w:sectPr"))
        if existing_section is not None:
            return

    target_section_properties = None
    for child in body_children[target_position:]:
        if child.tag == qn("w:p"):
            p_pr = child.find(qn("w:pPr"))
            if p_pr is not None:
                sect_pr = p_pr.find(qn("w:sectPr"))
                if sect_pr is not None:
                    target_section_properties = sect_pr
                    break
        elif child.tag == qn("w:sectPr"):
            target_section_properties = child
            break

    if target_section_properties is None:
        return

    if previous_p_pr is None:
        previous_p_pr = OxmlElement("w:pPr")
        previous_paragraph_xml.insert(0, previous_p_pr)

    previous_section = deepcopy(target_section_properties)

    # Do not carry a page-number restart back to the cover/index section.
    previous_page_number_type = previous_section.find(qn("w:pgNumType"))
    if previous_page_number_type is not None:
        previous_section.remove(previous_page_number_type)

    previous_p_pr.append(previous_section)


def _restart_financial_page_numbering(
    doc: DocxDocument,
    first_financial_title: Paragraph,
) -> None:
    """
    Restart Word's page numbering at 1 for the section containing
    'قائمة المركز المالي'.

    PAGEREF then returns report page numbers instead of cover/report pages.
    """
    _ensure_section_break_before_paragraph(
        doc,
        first_financial_title,
    )

    found_target = False
    target_section_properties = None

    for child in doc.element.body.iterchildren():
        if child is first_financial_title._p:
            found_target = True

        if not found_target:
            continue

        if child.tag == qn("w:p"):
            paragraph_properties = child.find(qn("w:pPr"))

            if paragraph_properties is not None:
                section_properties = paragraph_properties.find(qn("w:sectPr"))

                if section_properties is not None:
                    target_section_properties = section_properties
                    break

        elif child.tag == qn("w:sectPr"):
            target_section_properties = child
            break

    if target_section_properties is None:
        return

    page_number_type = target_section_properties.find(qn("w:pgNumType"))

    if page_number_type is None:
        page_number_type = OxmlElement("w:pgNumType")
        target_section_properties.append(page_number_type)

    page_number_type.set(qn("w:start"), "1")
    page_number_type.set(qn("w:fmt"), "decimal")

    # Continue page numbering through every later section. The template
    # originally restarts numbering again at the balance sheet, which would
    # otherwise make both the auditor report and balance sheet show page 1.
    passed_report_section = False

    for child in doc.element.body.iterchildren():
        section_properties = None

        if child.tag == qn("w:p"):
            paragraph_properties = child.find(qn("w:pPr"))
            if paragraph_properties is not None:
                section_properties = paragraph_properties.find(qn("w:sectPr"))
        elif child.tag == qn("w:sectPr"):
            section_properties = child

        if section_properties is None:
            continue

        if section_properties is target_section_properties:
            passed_report_section = True
            continue

        if not passed_report_section:
            continue

        later_page_number_type = section_properties.find(qn("w:pgNumType"))
        if later_page_number_type is not None:
            section_properties.remove(later_page_number_type)


def _enable_field_updates(doc: DocxDocument) -> None:
    """
    Ask Word to update PAGEREF fields when the generated DOCX is opened.
    """
    settings = doc.settings._element

    update_fields = settings.find(qn("w:updateFields"))

    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)

    update_fields.set(qn("w:val"), "true")


def update_page_two_and_index(
    doc: DocxDocument,
    form: Dict[str, str],
) -> None:
    """
    Main entry point.

    1. Fill page-2 heading from the Basic info / first-page values.
    2. Add the independent auditor report as page 1 in the index.
    3. Bookmark the report and financial-statement titles.
    4. Replace hard-coded index numbers with automatic PAGEREF fields.
    5. Make the notes page range automatic from first note to document end.
    """
    _update_page_two_heading(doc, form)

    auditor_index = _ensure_auditor_report_index_paragraph(doc)

    index_and_target_titles = [
        (
            "تقرير مدقق الحسابات المستقل",
            "تقرير مدقق الحسابات المستقل",
            "lw_auditor_report",
            "1",
        ),
        (
            "قائمة المركز المالي",
            "قائمة المركز المالي",
            "lw_balance_sheet",
            "4",
        ),
        (
            "قائمة الدخل الشامل",
            "قائمة الدخل الشامل",
            "lw_income_statement",
            "5",
        ),
        (
            "قائمة التغيرات في حقوق الملكية",
            "قائمة التغيرات في حقوق الملكية",
            "lw_equity_changes",
            "7",
        ),
        (
            "قائمة التدفقات النقدية",
            "قائمة التدفقات النقدية",
            "lw_cash_flows",
            "8",
        ),
    ]

    first_report_title = None

    for (
        index_title,
        target_title,
        bookmark_name,
        cached_page_number,
    ) in index_and_target_titles:
        if index_title == "تقرير مدقق الحسابات المستقل":
            index_paragraph = auditor_index
            target_paragraph = None

            for paragraph in doc.paragraphs:
                if _match_text(paragraph.text) == _match_text(target_title):
                    target_paragraph = paragraph
                    break
        else:
            index_paragraph = _find_index_paragraph(doc, index_title)
            target_paragraph = _find_exact_title_paragraph(doc, target_title)

        if index_paragraph is None or target_paragraph is None:
            continue

        if index_title == "تقرير مدقق الحسابات المستقل":
            first_report_title = target_paragraph

        _add_bookmark(doc, target_paragraph, bookmark_name)

        _replace_index_row(
            index_paragraph,
            index_title,
            [(bookmark_name, cached_page_number)],
        )

    notes_index = _find_index_paragraph(
        doc,
        "ايضاحات حول القوائم المالية",
    )
    notes_start = _find_notes_start(doc)
    notes_end = _find_last_body_paragraph(doc)

    if (
        notes_index is not None
        and notes_start is not None
        and notes_end is not None
    ):
        _add_bookmark(doc, notes_start, "lw_notes_start")
        _add_bookmark(doc, notes_end, "lw_notes_end")

        _replace_index_row(
            notes_index,
            "ايضاحـات حـول القوائم الماليـــة",
            [
                ("lw_notes_start", "5"),
                ("lw_notes_end", "21"),
            ],
        )

    if first_report_title is not None:
        _restart_financial_page_numbering(
            doc,
            first_report_title,
        )

    _enable_field_updates(doc)
